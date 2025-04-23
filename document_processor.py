"""
Module for processing various document formats
"""
import os
import re
import logging
import tempfile
import requests
from typing import List, Dict, Any, Optional, Tuple

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Chunk sizes
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 100

def extract_text_from_pdf(file_path: str, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[str]:
    """
    Extract text from a PDF file and split it into chunks
    
    Args:
        file_path: Path to the PDF file
        chunk_size: Maximum size of each text chunk
        chunk_overlap: Overlap between consecutive chunks
        
    Returns:
        List of text chunks extracted from the PDF
    """
    try:
        import pdfplumber
        
        # Open the PDF
        with pdfplumber.open(file_path) as pdf:
            # Extract text from each page
            all_text = ""
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    all_text += page_text + "\n\n"
            
            # Split into chunks with overlap
            if not all_text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")
                return []
            
            # Clean text: replace multiple newlines, spaces, etc.
            all_text = re.sub(r'\n{3,}', '\n\n', all_text)
            all_text = re.sub(r' {3,}', ' ', all_text)
            
            # Split text into chunks
            chunks = []
            start = 0
            text_length = len(all_text)
            
            while start < text_length:
                # Define end of chunk with overflow
                end = min(start + chunk_size, text_length)
                
                # If not at the end of the document and not at a natural break, try to find a good break point
                if end < text_length and all_text[end] not in ('.', '!', '?', '\n'):
                    # Look for natural break points (sentence ends, paragraphs)
                    break_point = all_text.rfind('. ', start, end)
                    if break_point == -1:
                        break_point = all_text.rfind('! ', start, end)
                    if break_point == -1:
                        break_point = all_text.rfind('? ', start, end)
                    if break_point == -1:
                        break_point = all_text.rfind('\n', start, end)
                    
                    # If found a good break point, use it
                    if break_point != -1:
                        end = break_point + 1
                
                # Add chunk
                chunks.append(all_text[start:end].strip())
                
                # Move start with overlap
                start = max(start + chunk_size - chunk_overlap, end - chunk_overlap)
            
            return chunks
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        return []

def extract_text_from_docx(file_path: str, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[str]:
    """
    Extract text from a DOCX file and split it into chunks
    
    Args:
        file_path: Path to the DOCX file
        chunk_size: Maximum size of each text chunk
        chunk_overlap: Overlap between consecutive chunks
        
    Returns:
        List of text chunks extracted from the DOCX
    """
    try:
        import docx
        
        # Open the DOCX
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        all_text = ""
        for para in doc.paragraphs:
            if para.text:
                all_text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ""
                for cell in row.cells:
                    if cell.text:
                        row_text += cell.text + " | "
                if row_text:
                    all_text += row_text.rstrip(" | ") + "\n"
                    
        # Split into chunks with overlap
        if not all_text.strip():
            logger.warning(f"No text extracted from DOCX: {file_path}")
            return []
        
        # Clean text: replace multiple newlines, spaces, etc.
        all_text = re.sub(r'\n{3,}', '\n\n', all_text)
        all_text = re.sub(r' {3,}', ' ', all_text)
        
        # Split text into chunks
        chunks = []
        start = 0
        text_length = len(all_text)
        
        while start < text_length:
            # Define end of chunk with overflow
            end = min(start + chunk_size, text_length)
            
            # If not at the end of the document and not at a natural break, try to find a good break point
            if end < text_length and all_text[end] not in ('.', '!', '?', '\n'):
                # Look for natural break points (sentence ends, paragraphs)
                break_point = all_text.rfind('. ', start, end)
                if break_point == -1:
                    break_point = all_text.rfind('! ', start, end)
                if break_point == -1:
                    break_point = all_text.rfind('? ', start, end)
                if break_point == -1:
                    break_point = all_text.rfind('\n', start, end)
                
                # If found a good break point, use it
                if break_point != -1:
                    end = break_point + 1
            
            # Add chunk
            chunks.append(all_text[start:end].strip())
            
            # Move start with overlap
            start = max(start + chunk_size - chunk_overlap, end - chunk_overlap)
        
        return chunks
    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return []

def extract_text_from_txt(file_path: str, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[str]:
    """
    Extract text from a text file (TXT or MD) and split it into chunks
    
    Args:
        file_path: Path to the text file
        chunk_size: Maximum size of each text chunk
        chunk_overlap: Overlap between consecutive chunks
        
    Returns:
        List of text chunks extracted from the text file
    """
    try:
        # Open and read the file
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            all_text = f.read()
        
        # Split into chunks with overlap
        if not all_text.strip():
            logger.warning(f"Empty text file: {file_path}")
            return []
        
        # Clean text: replace multiple newlines, spaces, etc.
        all_text = re.sub(r'\n{3,}', '\n\n', all_text)
        all_text = re.sub(r' {3,}', ' ', all_text)
        
        # Split text into chunks
        chunks = []
        start = 0
        text_length = len(all_text)
        
        while start < text_length:
            # Define end of chunk with overflow
            end = min(start + chunk_size, text_length)
            
            # If not at the end of the document and not at a natural break, try to find a good break point
            if end < text_length and all_text[end] not in ('.', '!', '?', '\n'):
                # Look for natural break points (sentence ends, paragraphs)
                break_point = all_text.rfind('. ', start, end)
                if break_point == -1:
                    break_point = all_text.rfind('! ', start, end)
                if break_point == -1:
                    break_point = all_text.rfind('? ', start, end)
                if break_point == -1:
                    break_point = all_text.rfind('\n', start, end)
                
                # If found a good break point, use it
                if break_point != -1:
                    end = break_point + 1
            
            # Add chunk
            chunks.append(all_text[start:end].strip())
            
            # Move start with overlap
            start = max(start + chunk_size - chunk_overlap, end - chunk_overlap)
        
        return chunks
    
    except Exception as e:
        logger.error(f"Error extracting text from text file {file_path}: {str(e)}")
        return []

def extract_text_from_xlsx(file_path: str, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP) -> List[str]:
    """
    Extract text from an Excel file and split it into chunks
    
    Args:
        file_path: Path to the Excel file
        chunk_size: Maximum size of each text chunk
        chunk_overlap: Overlap between consecutive chunks
        
    Returns:
        List of text chunks extracted from the Excel file
    """
    try:
        import pandas as pd
        
        # Read Excel file
        xl = pd.ExcelFile(file_path)
        
        # Process each sheet
        chunks = []
        for sheet_name in xl.sheet_names:
            # Read the sheet
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Clean NaN values
            df = df.fillna('')
            
            # Convert sheet to text
            sheet_text = f"Sheet: {sheet_name}\n\n"
            
            # Add column headers
            headers = " | ".join(str(col) for col in df.columns)
            sheet_text += headers + "\n" + "-" * len(headers) + "\n"
            
            # Add data rows
            for _, row in df.iterrows():
                row_text = " | ".join(str(val) for val in row)
                sheet_text += row_text + "\n"
            
            # Clean text
            sheet_text = re.sub(r' {3,}', ' ', sheet_text)
            
            # Create chunks by rows to keep context
            # Each chunk should contain the headers plus a reasonable number of rows
            rows = sheet_text.split('\n')
            
            if len(rows) <= 3:  # Headers only or empty sheet
                chunks.append(sheet_text)
                continue
            
            headers = rows[0] + "\n" + rows[1] + "\n" + rows[2] + "\n"
            current_chunk = headers
            
            for i in range(3, len(rows)):
                row = rows[i]
                
                # If adding this row would exceed chunk size, save current chunk and start new one
                if len(current_chunk) + len(row) > chunk_size and len(current_chunk) > len(headers):
                    chunks.append(current_chunk.strip())
                    current_chunk = headers + row + "\n"
                else:
                    current_chunk += row + "\n"
            
            # Add the last chunk if not empty
            if len(current_chunk) > len(headers):
                chunks.append(current_chunk.strip())
        
        logger.info(f"Extracted {len(chunks)} chunks from Excel: {file_path}")
        return chunks
    
    except Exception as e:
        logger.error(f"Error extracting text from Excel {file_path}: {str(e)}")
        return []

def process_document(file_path: str, file_id: str, file_name: str) -> Optional[List[str]]:
    """
    Process a document file based on its extension
    
    Args:
        file_path: Path to the file
        file_id: ID of the file
        file_name: Name of the file
        
    Returns:
        List of text chunks or None if processing failed
    """
    try:
        # Get file extension
        _, extension = os.path.splitext(file_name)
        extension = extension.lower()
        
        # Extract text based on file type
        if extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return extract_text_from_docx(file_path)
        elif extension in ['.txt', '.md']:
            return extract_text_from_txt(file_path)
        elif extension in ['.xlsx', '.xls']:
            return extract_text_from_xlsx(file_path)
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return None
    
    except Exception as e:
        logger.error(f"Error processing document {file_name}: {str(e)}")
        return None

def download_and_process_file(url: str, file_id: str, file_name: str, headers: Dict[str, str]) -> Optional[List[str]]:
    """
    Download a file from a URL and process it
    
    Args:
        url: URL to download the file
        file_id: ID of the file
        file_name: Name of the file
        headers: HTTP headers for the request
        
    Returns:
        List of text chunks or None if processing failed
    """
    try:
        # Create temp directory if it doesn't exist
        if not os.path.exists("temp"):
            os.makedirs("temp")
        
        file_path = f"temp/{file_name}"
        
        # Download the file
        response = requests.get(url, headers=headers)
        
        with open(file_path, "wb") as f:
            f.write(response.content)
        
        # Process the file
        text_chunks = process_document(file_path, file_id, file_name)
        
        # Clean up
        if os.path.exists(file_path):
            os.remove(file_path)
        
        return text_chunks
    
    except Exception as e:
        logger.error(f"Error downloading and processing file {file_name}: {str(e)}")
        return None